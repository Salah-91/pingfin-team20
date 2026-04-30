"""
Voegt belangrijke screenshots in op de juiste plekken in PingFin_Eindverslag.docx.
Output: PingFin_Eindverslag_met_screenshots.docx
"""
import os
from copy import deepcopy
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SCREENSHOTS_DIR = r"C:\temp\screenshots\internationaalweek screenshots"
INPUT_DOC = r"C:\Users\Salah\Downloads\PingFin_Eindverslag.docx"
OUTPUT_DOC = r"C:\Users\Salah\Downloads\PingFin_Eindverslag_met_screenshots.docx"


def s(name):
    return os.path.join(SCREENSHOTS_DIR, f"Schermafbeelding {name}.png")


# Screenshot → (heading prefix to find, caption, width inches)
INSERTIONS = [
    # 3. Algemeen Messaging Schema
    ("3. Algemeen Messaging Schema",
     s("2026-04-29 094917"),
     "Figuur 1 — Algemeen messaging schema (manual): PO_NEW → OB → CB → BB → ACK terug.",
     6.0),

    # 3.1 Use Cases — onze eigen flowchart
    ("3.1 Use Cases",
     s("2026-04-27 120001"),
     "Figuur 2 — Use cases UC1-UC5 met legende. Groen = validation passed, rood = validation failed, blauw = message, oranje = account processing.",
     6.5),

    # 4.1 Globale Structuur — onze repo / docker
    ("4.1 Globale Structuur",
     s("2026-04-29 084805"),
     "Figuur 3 — Docker Compose opstart van beide bank-services + gedeelde MySQL.",
     6.0),

    # 5. Database — MySQL data
    ("5. Database",
     s("2026-04-27 230533"),
     "Figuur 4 — MySQL Workbench met alle tabellen van pingfin_b1 en pingfin_b2.",
     5.5),

    # 7. GUI — Dashboard PO Aanmaken (Bank 1)
    ("7. GUI",
     s("2026-04-29 135105"),
     "Figuur 5 — PingFin Bank Dashboard (Bank 1, CEKVBE88) — tab \"PO Aanmaken\" met manuele PO formulier en succesvolle PENDING-status (code 2000).",
     6.5),

    # 7.1 Tabs — Dashboard met PO_OUT lijst
    ("7.1 Tabs",
     s("2026-04-29 115638"),
     "Figuur 6 — Bank 2 GUI met 5 succesvol verwerkte PO's (allemaal code 2000) zichtbaar in het log-paneel.",
     6.5),

    # 7.2 Live Updates — toast notifications
    ("7.2 Live Updates",
     s("2026-04-29 114848"),
     "Figuur 7 — Live updates: PO's met error code 4101 (ACCOUNT_UNKNOWN) worden meteen rood weergegeven in het log-paneel.",
     6.5),

    # 9. Beveiliging — 8 verdedigingslagen
    ("9. Beveiliging",
     s("2026-04-30 141038"),
     "Figuur 8 — De 8 verdedigingslagen (defense in depth) zoals gepresenteerd in de eindpresentatie.",
     6.0),

    # 10.2 Productie-Statistieken
    ("10.2 Productie-Statistieken",
     s("2026-04-30 122509"),
     "Figuur 9 — Live productie-statistieken: 1066+ ACKs verwerkt, 51 banken bij CB, 58/58 tests groen, 24/7 online.",
     6.0),

    # 11.2 Cloud Deployment (Railway)
    ("11.2 Cloud Deployment (Railway)",
     s("2026-04-30 110353"),
     "Figuur 10 — Railway settings van pingfin-team20-bank1 met Root Directory ingesteld op 'Bank 1' voor correcte Dockerfile-pickup.",
     6.0),

    # Dag 1 — Trello task setup
    ("Dag 1 — Analyse & Planning",
     s("2026-04-27 121509"),
     "Figuur 11 — Trello task management voor dag 1: Task B (Repository setup) afgerond met 86% checkbox-progressie.",
     6.0),

    # 13.2 Wat ging minder — push error
    ("13.2 Wat ging minder",
     s("2026-04-27 135302"),
     "Figuur 12 — \"Failed to push\" — branch protection rules vergden eerst een pull request in plaats van direct push naar main. Heeft ons geleerd om de protection-rules vroeger uit te denken.",
     5.5),
]


def find_heading_paragraph(doc, prefix):
    """Vind paragraaf met heading-style waarvan tekst begint met prefix."""
    for i, p in enumerate(doc.paragraphs):
        if "Heading" in p.style.name and p.text.strip().startswith(prefix):
            return i, p
    return None, None


def insert_image_after(doc, anchor_para, img_path, caption, width_inches):
    """Voeg een paragraaf met afbeelding + caption toe ná de anchor-paragraaf."""
    if not os.path.exists(img_path):
        print(f"    [X] Image niet gevonden: {os.path.basename(img_path)}")
        return False

    # Maak nieuwe paragraaf NA de anchor — gebruik low-level XML manipulation
    new_p = doc.paragraphs[0].insert_paragraph_before("")  # placeholder
    # Beter: maak paragraaf in body en plaats hem ná anchor via XML
    from docx.oxml.ns import qn

    # Nieuwe paragraaf voor afbeelding
    img_p = doc.add_paragraph()
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_p.add_run()
    run.add_picture(img_path, width=Inches(width_inches))

    # Caption-paragraaf
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap_p.add_run(caption)
    cap_run.italic = True
    cap_run.font.size = Pt(10)
    cap_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Verplaats img_p en cap_p naar direct ná anchor_para
    anchor_xml = anchor_para._element
    img_xml = img_p._element
    cap_xml = cap_p._element
    # Remove from current position
    img_xml.getparent().remove(img_xml)
    cap_xml.getparent().remove(cap_xml)
    # Insert after anchor: anchor → next sibling becomes img, then caption
    anchor_xml.addnext(cap_xml)   # caption goes second-to-anchor
    anchor_xml.addnext(img_xml)   # then image goes between anchor and caption

    # Remove the placeholder we created at top
    placeholder_xml = new_p._element
    if placeholder_xml.getparent() is not None:
        placeholder_xml.getparent().remove(placeholder_xml)

    return True


def main():
    print(f"Loading {INPUT_DOC}")
    doc = Document(INPUT_DOC)
    print(f"  Paragraphs: {len(doc.paragraphs)}")

    inserted = 0
    skipped = 0
    for prefix, img_path, caption, width in INSERTIONS:
        idx, anchor = find_heading_paragraph(doc, prefix)
        if anchor is None:
            print(f"  [?] Heading niet gevonden: {prefix}")
            skipped += 1
            continue
        ok = insert_image_after(doc, anchor, img_path, caption, width)
        if ok:
            print(f"  [OK] Inserted after \"{prefix}\": {os.path.basename(img_path)}")
            inserted += 1
        else:
            skipped += 1

    print(f"\nSaving to {OUTPUT_DOC}")
    doc.save(OUTPUT_DOC)
    print(f"Done. Inserted: {inserted} | Skipped: {skipped}")


if __name__ == "__main__":
    main()
